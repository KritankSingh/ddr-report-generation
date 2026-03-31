"""Generate DOCX report from structured DDR data and images."""
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def save_image_bytes(image_bytes, ext="png"):
    """Save image bytes to a temporary file and return the path."""
    temp_file = tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False)
    temp_file.write(image_bytes)
    temp_file.close()
    return temp_file.name


def generate_docx(ddr_data, inspection_images, thermal_images, output_path):
    """
    Generate a DOCX report from DDR data and images.

    Args:
        ddr_data: dict with structured DDR content
        inspection_images: list of image dicts from inspection PDF
        thermal_images: list of image dicts from thermal PDF
        output_path: str or Path for output .docx file
    """
    doc = Document()
    output_path = Path(output_path)

    # Title
    doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    doc.add_paragraph("")

    # Section 1: Property Issue Summary
    doc.add_heading("1. Property Issue Summary", level=1)
    summary = ddr_data.get("property_issue_summary", "Not Available")
    doc.add_paragraph(summary if summary else "Not Available")
    doc.add_paragraph("")

    # Section 2: Area-wise Observations
    doc.add_heading("2. Area-wise Observations", level=1)
    area_obs = ddr_data.get("area_observations", [])
    if area_obs:
        for obs in area_obs:
            area = obs.get("area", "Unknown Area")
            findings = obs.get("findings", "Not Available")
            severity = obs.get("severity", "Not Available")
            severity_reason = obs.get("severity_reason", "")

            # Sub-heading for area
            doc.add_heading(f"  • {area}", level=2)
            doc.add_paragraph(f"Findings: {findings}")
            doc.add_paragraph(f"Severity: {severity} — {severity_reason}")

            # Add thermal images if referenced
            thermal_ref = obs.get("thermal_image_references", "None available")
            thermal_indices = obs.get("thermal_image_indices", [])

            if thermal_indices and thermal_images:
                doc.add_paragraph(f"Thermal Data: {thermal_ref}")
                # Add each relevant thermal image for this area
                images_added = 0
                for img_idx in thermal_indices:
                    if 0 <= img_idx < len(thermal_images):
                        try:
                            thermal_img = thermal_images[img_idx]
                            img_path = save_image_bytes(thermal_img['bytes'], thermal_img.get('ext', 'png'))
                            doc.add_picture(img_path, width=Inches(4))
                            images_added += 1
                        except Exception as e:
                            print(f"Warning: Could not embed thermal image {img_idx}: {e}")

                if images_added == 0:
                    doc.add_paragraph("Image Not Available")
            else:
                doc.add_paragraph("Image Not Available")

            doc.add_paragraph("")
    else:
        doc.add_paragraph("Not Available")
        doc.add_paragraph("")

    # Section 3: Probable Root Cause
    doc.add_heading("3. Probable Root Cause", level=1)
    root_cause = ddr_data.get("probable_root_causes", "Not Available")
    doc.add_paragraph(root_cause if root_cause else "Not Available")
    doc.add_paragraph("")

    # Section 4: Severity Assessment
    doc.add_heading("4. Severity Assessment", level=1)
    severity_info = ddr_data.get("severity_assessment", {})
    overall = severity_info.get("overall_level", "Not Available")
    reasoning = severity_info.get("reasoning", "")
    doc.add_paragraph(f"Overall Severity Level: {overall}")
    if reasoning:
        doc.add_paragraph(f"Reasoning: {reasoning}")
    doc.add_paragraph("")

    # Section 5: Recommended Actions
    doc.add_heading("5. Recommended Actions", level=1)
    actions = ddr_data.get("recommended_actions", [])
    if actions:
        for action in actions:
            doc.add_paragraph(action, style='List Bullet')
    else:
        doc.add_paragraph("Not Available")
    doc.add_paragraph("")

    # Section 6: Additional Notes
    doc.add_heading("6. Additional Notes", level=1)
    notes = ddr_data.get("additional_notes", "Not Available")
    doc.add_paragraph(notes if notes else "Not Available")
    doc.add_paragraph("")

    # Section 7: Missing or Unclear Information
    doc.add_heading("7. Missing or Unclear Information", level=1)
    missing = ddr_data.get("missing_or_unclear_info", [])
    if missing:
        for item in missing:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("All information available")
    doc.add_paragraph("")

    # Save the document
    doc.save(str(output_path))
    print(f"Report generated: {output_path}")
